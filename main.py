import json
import string
import re
from pprint import pprint
import matplotlib.pyplot as plt
import numpy as np
import seaborn as sns
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from tabulate import tabulate
from collections import Counter
from operator import itemgetter
import os

path = os.path.join(os.getcwd(), "res")

if not os.path.exists(path):
    os.mkdir(path)


def create_table(name, file):                                            # функция обработки таблиц: переводится pandas.DataFrame в таблицу ворд
    table = file.add_table(rows=len(name) + 1, cols=len(name.columns))

    for i, column_name in enumerate(name.columns):
        table.cell(0, i).text = column_name

    for row in range(len(name)):
        for col in range(len(name.columns)):
            table.cell(row + 1, col).text = str(name.iloc[row, col])


def at(file, string):                                                   # функция добавления надписи в docx
    file.add_paragraph(string)


with open('result.json', 'r', encoding='utf-8') as f:                   # открытие json-файла
    text = json.load(f)
data = []


for i in text['messages']:                                              #цикл-обработчик-парсер информации из json
    if 'sticker_emoji' in i:                                            #выделение стикеров
        data.append([i['id'], i['date'], i['media_type'], i['from'], i['sticker_emoji']])
    if 'action' in i:                                                   #звонки
        data.append([i['id'], i['date'], i['type'] + '_' + i['action'], i['actor']])
        continue
    if 'mime_type' in i and 'application' in i['mime_type']:            #пропуск различных медиаданных, если они прилетают
        continue
    if 'location_information' in i:                                     #пропуск геоданных
        continue
    if 'thumbnail' in i and 'media_type' not in i:                      #пропуск файлов
        continue
    if 'photo' in i:                                                    #обработка фотографий, как текстовых данных
        data.append([i['id'], i['date'], i['type'] + "_photo", i['from']])
        continue
    if "poll" in i:                                                     #пропуск голосований
        continue                                                        #далее идут обработчики различных перекликающихся форматов
                                                                        #например, аудио/видео сообщения, текстовых выражений: стикеров, массивов строк, и, в целом, сообщений
    if len(i['text_entities']) == 1 and i['text_entities'][0]['type'] == 'plain':
        data.append([i['id'], i['date'], i['type'], i['from'], i['text']])
        continue
    if (len(i['text_entities']) == 2 or len(i['text_entities']) == 0) and 'media_type' not in i:
        continue;
    if len(i['text_entities']) == 2 or (len(i['text_entities']) == 0 and (i['media_type'] == 'video_message' or (i['media_type'] == 'voice_message' and 'duration_seconds' in i))):
        title = i['title'] if ['media_type'] == 'audio_file' else np.nan
        data.append([i['id'], i['date'], i['media_type'], i['from'], title, i['duration_seconds']])
        continue
    if len(i['text_entities']) == 0 and i['media_type'] == 'sticker':
        if 'sticker_emoji' not in i:
            continue
        data.append([i['id'], i['date'], i['media_type'], i['from'], i['sticker_emoji']])
        continue
    if len(i['text_entities']) == 1 and isinstance(i['text'], list):
        data.append([i['id'], i['date'], i['text'][0]['type'], i['from'], i['text'][0]['text']])
        continue
    if 'performer' not in i or 'title' not in i:
        # if i['id'] == 178677: print(1)
        continue
    if len(i['text_entities']) == 0 and i['media_type'] == 'audio_file':
        dur = i['duration_seconds'] if 'duration_seconds' in i else np.nan
        data.append([i['id'], i['date'], i['media_type'], i['from'], i['performer'] + '_' + i['title'], dur])
        continue

ll = 0
for item in data:
   if len(item) > ll:
       ll = len(item)
if ll == 5:
    for item in data:
        item.append(0)
cols = ['id', 'date', 'type', 'from', 'text', 'seconds']              #формируем колонки
df = pd.DataFrame(data, columns=cols)
df.to_csv("res.csv")                                                  #конвертируем получившийся dataframe в сsv-файл, для различных нужд и сохраняем в файл в корне

data = pd.read_csv('res.csv', index_col=0)
                                                                      #избавляемся от стикеров в названии собеседников, убираем пробелы при необходимости и делаем _ между пробелами
data['from'] = [re.sub('[^\x00-\x7Fа-яА-Я]', '', i) for i in data['from']]
data['from'] = [i[:-1] if (i[-1] == ' ' and len(i) > 1) else i for i in data['from']]
data['from'] = [i.replace(' ', '_') for i in data['from']]
                                                                     #создание массивов из 2 элементов - это собеседники и избавляемся от NaN
mass_from = data['from'].unique()
data = data.fillna(0)
                                                                     #создаем первый пай, основанный на вовлеченности по сообщениям
colors = sns.color_palette('pastel')[0:2]
wrt = data['from'].value_counts()
plt.figure(figsize=(7, 7))
plt.pie(wrt, colors=colors, autopct="%.2f%%")
plt.title('Involvement diagram', fontsize=20)
plt.legend(labels=wrt.index)
plt.savefig("res/involvment.png")
                                                                    #создаем наследника класса Document, куда непосредственно будет все записываться
file = Document()
heading = file.add_heading(mass_from[0] + " и " + mass_from[1] + " анализ переписки", level=1)
                                                                    #работа с дизайном
run = heading.runs[0]
font = run.font
font.name = 'Times New Roman'
font.size = Pt(16)
font.color.rgb = RGBColor(0, 0, 0)
                                                                    #добавление диаграммы по вовлеченности
file.add_picture('res/involvment.png', width=Inches(6), height=Inches(6))

wrt = pd.DataFrame(wrt).reset_index()
wrt = wrt.rename(columns={
    'index':'name',
    'from':'count'
})
                                                                    #добавление первой таблицы в файл
at(file, "Количество сообщений:")
create_table(wrt, file)
                                                                    #создается словарь со значением месяц-цифра для форматирования времени
uio = data['date'].apply(lambda x: str(x)[:-9].replace('-', '.'))
keys = ['янв', 'фев', 'мар', 'апр', 'май', 'июн', 'июл', 'авг', 'сен', 'окт', 'ноя', 'дек']
val = ["01", "02", "03", "04", "05", "06", "07", "08", "09", "10", "11", "12"]
diction = dict(zip(val, keys))
data['mon_day_hou'] = uio                                           #колонка mon_day_hoy - форматированная дата
times = data[['mon_day_hou', 'from']]

jio = times.value_counts().to_frame().sort_values('mon_day_hou', ascending=True).reset_index()
jio = jio.rename(columns={
    0:"counts"
})
nw = times['mon_day_hou'].value_counts().to_frame().sort_index().reset_index()
nw = nw.rename(columns={
    "index":"mon_day_hou",
    "mon_day_hou":"counts"
})
jio_d = jio[jio['from'] == mass_from[0]]                            #выполняется личный подсчет по каждому собеседнику по дню
jio_m = jio[jio['from'] == mass_from[1]]
st = nw.merge(jio_m, how="left", left_on="mon_day_hou", right_on="mon_day_hou").fillna(0)    #данные сливаются, чтобы создать общую таблицу
st = st.merge(jio_d, how="outer", left_on="mon_day_hou", right_on="mon_day_hou").fillna(0)
st.head()

st = st.rename(columns={
    "counts_x":"all_count",
    "counts_y": mass_from[0]+"_count",
    "counts": mass_from[1]+"_count"
})
                                                                        #непосредственно форматирование DataFrame и даты
st = st.drop(['from_x', 'from_y'], axis=1)
st[[mass_from[0]+"_count", mass_from[1]+"_count"]] = st[[mass_from[0]+"_count", mass_from[1]+"_count"]].astype('int')
st['mon_day_hou'] = st['mon_day_hou'].apply(lambda x: x[8:] + " " + diction[str(x[5:7])] + " " + x[:4])
                                                                        #создание бара по каждому собеседнику по каждому дню
fig, ax = plt.subplots(figsize=(14, 7))
x = np.arange(len(st['mon_day_hou']))
width = 0.2

rects1 = ax.bar(x - width/2, st[mass_from[1]+"_count"], width, label=mass_from[1])
rects2 = ax.bar(x + width/2, st[mass_from[0]+"_count"], width, label=mass_from[0])
ax.set_label('message count')
ax.set_title('Stat for messages/day')
ax.set_xticks(x)
ax.set_xticklabels(st['mon_day_hou'])
ax.legend()
fig.tight_layout()
plt.xticks(rotation=90)
plt.savefig("res/statForMessages.png")
                                                                        #загрузка получившейся картинки в файл
file.add_picture('res/statForMessages.png', width=Inches(7), height=Inches(4))
                                                                        #получение статистики по максимальному, минимальному и среднему количеству сообщений по дням через агрегацию
at(file, "Количество сообщений за день (в столбец: среднее, максимальное, минимальное):")
mmm_stat = st[['all_count', mass_from[0]+'_count', mass_from[1]+'_count']].agg({'max', 'min', 'mean'}).round(2).reset_index()
mmm_stat = mmm_stat.rename(columns={
    "index":"velichina"
})
create_table(mmm_stat, file)

                                                                        #подсчет самых поплярных и менее популярных дней по сообщениям
file.add_paragraph('Наиболее популярные по переписке дни:')
mcd = st.sort_values('all_count', ascending=False).reset_index(drop=True).loc[0:4]
create_table(mcd, file)

at(file, "Наименее популярные по переписке дни:")
msd = st.sort_values('all_count', ascending=True).reset_index(drop=True).loc[0:4]
create_table(msd, file)

ui = times['mon_day_hou'].value_counts().sort_index().to_frame().reset_index().rename(columns={'index':'mon_day_hou', 'mon_day_hou':'count'})
ui['mon_day_hou'] = ui['mon_day_hou'].apply(lambda x: x[8:] + " " + diction[str(x[5:7])] + " " + x[:4])
                                                                        #создание графика развития переписки
plt.figure(figsize=(16, 8))
plt.grid()
plt.title('Conversation development')
plt.xticks(rotation=90)
plt.scatter(ui['mon_day_hou'], ui['count'], marker='o', color='black')
plt.plot(ui['mon_day_hou'], ui['count'], label='conversation', lw=3, mec='b', mew=2, ms=10, color=sns.color_palette('dark')[0])
plt.legend()
plt.savefig("res/conversationDev.png")
file.add_picture('res/conversationDev.png', width=Inches(7), height=Inches(4))
                                                                        #выделение только сообщений и стикеров в одну таблицу для анализа
messages = data[['type', 'from', 'text', 'mon_day_hou']]
messages = messages.loc[(messages['type'] == 'message') | (messages['type'] == 'sticker')]
pd.options.mode.chained_assignment = None
messages['text'] = messages['text'].astype('str')
messages['text'] = messages['text'].apply(lambda x: x + ' ')
                                                                        #создание массива из слов по каждому сообщению и его форматирование
hj = messages.groupby('mon_day_hou').agg({'text':'sum'}).reset_index()
hj['text'] = hj['text'].apply(lambda x: x.split())
hj['text'] = hj['text'].apply(lambda x: [re.sub('[!@#%&(){};:,./<>?\|`~=_+]', '', i) for i in x])
hj['words'] = hj['text'].apply(lambda x: len(x))
                                                                        #создание графика по развитию диплога именно по словам
plt.figure(figsize=(16, 8))
plt.grid()
plt.title('Total words development')
plt.xticks(rotation=90)
plt.scatter(hj['mon_day_hou'], hj['words'], marker='o', color='black')
plt.plot(hj['mon_day_hou'], hj['words'], label='total words', lw=3, mec='b', mew=2, ms=10, color=sns.color_palette('dark')[4])
plt.legend()
plt.savefig("res/totalWordsDev.png")
file.add_picture('res/totalWordsDev.png', width=Inches(7), height=Inches(4))
                                                                        #для выделение 50 самых поплярных слов по переписке и потом по каждому участнику были выбраны слова, у которвх больше 4 символов для лучшего осмысления и анализа
                                                                        #основные подсчеты выполняются через Counter
hj['unique_words'] = hj['text'].apply(lambda x: len(set(x)))
some = hj['text'].apply(lambda x: dict(sorted(Counter(x).items(), key=itemgetter(1))))
an_tty = hj['text'].apply(lambda x: [i for i in x if (len(i) > 4 and len(''.join(set(i))) > 2) or i == 'я'])
tty = hj['text'].apply(lambda x: [i for i in x if (len(i) > 4 and len(''.join(set(i))) > 2) or i == 'я'])
words = pd.Series(tty.agg('sum')).value_counts().to_frame().reset_index().rename(columns={'index':'word', 0:'count'}).loc[0:50]

at(file, "50 самых популярных слов в переписке")
create_table(words, file)
                                                                        #форматировнание и подсчет статистики по словам, уникальным словам и проценту уникальных слов
first = messages[messages['from'] == mass_from[1]]
second = messages[messages['from'] == mass_from[0]]
first = first.groupby('mon_day_hou').agg({'text':'sum'}).reset_index()
second = second.groupby('mon_day_hou').agg({'text':'sum'}).reset_index()
first['text'] = first['text'].apply(lambda x: x.split()).apply(lambda x: [re.sub('[!@#%&(){};:,./<>?\|`~=_+]', '', i) for i in x])
second['text'] = second['text'].apply(lambda x: x.split()).apply(lambda x: [re.sub('[!@#%&(){};:,./<>?\|`~=_+]', '', i) for i in x])
first['words'] = first['text'].apply(lambda x: len(x))
second['words'] = second['text'].apply(lambda x: len(x))
ggg_1 = pd.DataFrame({
    'name': mass_from[1],
    "Words": first['words'].sum(),
    'Unique Words': len(set(first['text'].sum())),
    'Percentele of unique' : str(round(len(set(first['text'].sum()))/first['words'].sum()*100, 2))+'%'}, index=[0])
ggg_2 = pd.DataFrame({
    'name': mass_from[0],
    "Words": second['words'].sum(),
    'Unique Words': len(set(second['text'].sum())),
    'Percentele of unique' : str(round(len(set(second['text'].sum()))/second['words'].sum()*100, 2))+'%'}, index=[1])
ggg = pd.concat([ggg_1, ggg_2])
at(file, "Статистика по словам:")
create_table(ggg, file)
                                                                        #выделение 50 самых популярных слов по каждому участнику
tty_f = first['text'].apply(lambda x: [i for i in x if (len(i) > 4 and len(''.join(set(i))) > 2) or i == 'я'])
tty_s = second['text'].apply(lambda x: [i for i in x if (len(i) > 4 and len(''.join(set(i))) > 2) or i == 'я'])

words_1 = pd.Series(tty_f.agg('sum')).value_counts().to_frame().reset_index().rename(columns={'index':'words_'+mass_from[1], 0:'count'}).loc[0:50]
at(file, "50 самых популярных слов " + mass_from[1] + ":")
create_table(words_1, file)
                                                                        #
words_2 = pd.Series(tty_s.agg('sum')).value_counts().to_frame().reset_index().rename(columns={'index':'words_'+mass_from[0], 0:'count'}).loc[0:50]
at(file, "50 самых популярных слов " + mass_from[0] + ":")
create_table(words_2, file)
                                                                        #диаграмма различных типом медиа в процентном соотношении
osob = data[data['type'] != 'message']
dfg = osob['type'].value_counts()
plt.figure(figsize=(10, 10))
plt.title("Difference of media", fontsize=20)
color = sns.color_palette('pastel')[0:7]
plt.pie(dfg, colors=color, autopct="%.2f%%")
plt.legend(bbox_to_anchor=(0.7, 0.7),labels=dfg.index)
plt.savefig("res/diffOfMedia.png")
file.add_picture('res/diffOfMedia.png', width=Inches(5), height=Inches(5))
                                                                        #диаграмма участников по видеосообщениям и среднее время видеосообщения
vid = osob[osob['seconds'] != 0]
vid = vid[vid['type'] == "video_message"]
v = vid['from'].value_counts()
plt.figure(figsize=(7, 7))
plt.title('Video messages diagram', fontsize=20)
color = sns.color_palette('pastel')[2:4]
plt.pie(v, colors=color, autopct="%.2f%%")
plt.legend(labels=v.index)
plt.savefig("res/videoMessages.png")
file.add_picture('res/videoMessages.png', width=Inches(5), height=Inches(5))
at(file, 'Среднее время видеосообщения: '+ str(round(vid['seconds'].mean(), 2)))
                                                                        #статистика участников по видеосообщениям
v_stat = vid.groupby('from')['seconds'].agg({'sum', 'mean', 'count'}).reset_index()
at(file, "Средняя статистика по видеосообщениям в секундах")
create_table(v_stat, file)
                                                                        #диаграмма участников по аудиосообщениям
voice = osob[osob['type'] == 'voice_message'].reset_index(drop=True)
vo = voice['from'].value_counts()
plt.figure(figsize=(7, 7))
plt.title('Voice messages diagram', fontsize=20)
color = sns.color_palette('pastel')[0:len(vo)]
plt.pie(vo, colors=color, autopct="%.2f%%")
plt.legend(labels=vo.index)
plt.savefig("res/voice.png")
file.add_picture('res/voice.png', width=Inches(5), height=Inches(5))
                                                                        # статистика участников по аудиосообщениям
at(file, "Средняя статистика по голосовым сообщениям в секундах")
v_df = voice.groupby('from')['seconds'].agg({'sum', 'mean', 'count'}).reset_index()
create_table(v_df, file)
                                                                        # статистика участников по стикерам
stick = osob[osob['type'] == 'sticker'].reset_index(drop=True)
s = stick['from'].value_counts()
plt.figure(figsize=(7, 7))
plt.title('Stickers diagram', fontsize=20)
color = sns.color_palette('pastel')[0:len(s)]
plt.pie(s, colors=color, autopct="%.2f%%")
plt.legend(labels=s.index)
plt.savefig("res/stickers.png")
file.add_picture('res/stickers.png', width=Inches(5), height=Inches(5))
                                                                        # статистика участников по стикерам
s_df = stick.groupby('from')['seconds'].agg({'count'}).reset_index()
at(file, "Средняя статистика по количеству стикеров:")
create_table(s_df, file)
                                                                        #сохранение файла в корень
file.save('analize.docx')
